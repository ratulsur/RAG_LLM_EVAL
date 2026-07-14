"""
PATHOLOGY (b): real content hides in TEXT BOXES, HEADERS/FOOTERS, and FOOTNOTES
-- parts that naive body extraction never visits.

WHY NAIVE EXTRACTION FAILS
--------------------------
python-docx's `document.paragraphs` walks only the main body. It does not read:
  - <w:txbxContent> inside a <w:pict>/drawing (text boxes, callouts, sidebars)
  - header1.xml / footer1.xml (running text, sometimes real clauses/definitions)
  - footnotes.xml (the substance of legal/academic docs often lives here)
On our fixture the sidebar "escrow release requires dual sign-off", the header
definition, and the liability-cap footnote all disappear. A RAG answer about
escrow or the liability cap would then be a confident hallucination.

THE PRODUCTION FIX
------------------
Enumerate the ZIP parts and pull text from each hiding place explicitly, tagging
provenance so chunks stay attributable:
  textbox -> word/document.xml  //w:txbxContent
  header  -> word/header*.xml
  footer  -> word/footer*.xml
  footnote-> word/footnotes.xml (skip the separator pseudo-footnotes id<=0)

SENIOR TRADEOFF
---------------
Headers/footers are double-edged: a running "CONFIDENTIAL" banner is furniture
(drop, like the PDF case), but a header that carries a defined term or a clause
is content. Heuristic: if a header/footer string recurs identically across
sections it's likely furniture; if it's unique or long it's content. We SURFACE
it with a provenance tag and let the chunker's furniture filter decide, rather
than dropping blind. Footnote text should be chunked NEAR its reference, not
appended at the end, or you break the local context.

Maps to Universal Document Ingestor: the "recover-everything" pass that
guarantees no addressable content is left in a part the body walker skipped.
"""
from __future__ import annotations

import os
import zipfile
from dataclasses import dataclass
from typing import List

from lxml import etree

HERE = os.path.dirname(os.path.abspath(__file__))
FIX = os.path.join(HERE, "fixtures", "dirty.docx")
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _qn(t: str) -> str:
    return f"{{{W}}}{t}"


@dataclass
class Hidden:
    kind: str      # textbox | header | footer | footnote
    source: str    # zip member
    text: str


class HiddenContentRecoverer:
    def __init__(self, path: str):
        self.path = path

    def _members(self):
        with zipfile.ZipFile(self.path) as z:
            return z.namelist()

    def _root(self, member: str):
        with zipfile.ZipFile(self.path) as z:
            return etree.fromstring(z.read(member))

    def _text_of(self, el) -> str:
        return "".join(t.text or "" for t in el.iter(_qn("t"))).strip()

    def recover(self) -> List[Hidden]:
        found: List[Hidden] = []
        members = self._members()

        # text boxes live in the main document part
        if "word/document.xml" in members:
            doc = self._root("word/document.xml")
            for tb in doc.iter(_qn("txbxContent")):
                txt = self._text_of(tb)
                if txt:
                    found.append(Hidden("textbox", "word/document.xml", txt))

        # headers / footers
        for m in members:
            base = m.split("/")[-1]
            if base.startswith("header") and base.endswith(".xml"):
                txt = self._text_of(self._root(m))
                if txt:
                    found.append(Hidden("header", m, txt))
            elif base.startswith("footer") and base.endswith(".xml"):
                txt = self._text_of(self._root(m))
                if txt:
                    found.append(Hidden("footer", m, txt))

        # footnotes (skip separator/continuation pseudo-notes with id <= 0)
        if "word/footnotes.xml" in members:
            fr = self._root("word/footnotes.xml")
            for fn in fr.iter(_qn("footnote")):
                try:
                    fid = int(fn.get(_qn("id")))
                except (TypeError, ValueError):
                    continue
                if fid <= 0:
                    continue
                txt = self._text_of(fn)
                if txt:
                    found.append(Hidden("footnote", "word/footnotes.xml", txt))
        return found


def naive_body_only(path: str) -> str:
    import docx
    return "\n".join(p.text for p in docx.Document(path).paragraphs if p.text.strip())


if __name__ == "__main__":
    if not os.path.exists(FIX):
        from fixtures.generate import build_all
        build_all()

    body = naive_body_only(FIX)
    print("=" * 70)
    print("NAIVE body-only extraction MISSES: textbox, header, footer, footnote")
    print("=" * 70)
    for probe in ["escrow", "MASTER SERVICES", "trailing twelve"]:
        print(f"  '{probe}' present in naive body? {probe.lower() in body.lower()}")

    print("\n" + "=" * 70)
    print("RECOVERED hidden content (with provenance tags):")
    print("=" * 70)
    for h in HiddenContentRecoverer(FIX).recover():
        print(f"  [{h.kind:8}] ({h.source}) {h.text}")
